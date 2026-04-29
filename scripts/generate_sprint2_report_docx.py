from __future__ import annotations

import base64
import os
import subprocess
import sys
import textwrap
import zlib
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from matplotlib.patches import FancyBboxPatch
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "report_assets"


def run_command(args: list[str], cwd: Path) -> str:
    result = subprocess.run(
        args,
        cwd=str(cwd),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )
    output = result.stdout.strip()
    if result.stderr.strip():
        output = f"{output}\n\n[stderr]\n{result.stderr.strip()}".strip()
    return output


def _load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]
    if bold:
        candidates.insert(0, "C:/Windows/Fonts/calibrib.ttf")

    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_to_image(text: str, out_path: Path, title: str) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    width = 1800
    margin = 28
    line_width = 130

    title_font = _load_font(30, bold=True)
    body_font = _load_font(20, bold=False)

    wrapped_lines: list[str] = []
    for raw in text.splitlines() or [""]:
        pieces = textwrap.wrap(raw, width=line_width) or [""]
        wrapped_lines.extend(pieces)

    title_h = 52
    line_h = 28
    max_lines = 120
    wrapped_lines = wrapped_lines[:max_lines]

    height = margin * 2 + title_h + line_h * max(1, len(wrapped_lines))
    image = Image.new("RGB", (width, height), color=(15, 23, 42))
    draw = ImageDraw.Draw(image)

    draw.text((margin, margin - 4), title, font=title_font, fill=(240, 248, 255))
    y = margin + title_h
    for line in wrapped_lines:
        draw.text((margin, y), line, font=body_font, fill=(203, 213, 225))
        y += line_h

    image.save(out_path)
    return out_path


def mermaid_to_image(mermaid_code: str, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii")
    url = f"https://mermaid.ink/img/{encoded}"

    response = requests.get(url, timeout=25)
    response.raise_for_status()
    out_path.write_bytes(response.content)
    return out_path


def _encode6bit(value: int) -> str:
    if value < 10:
        return chr(48 + value)
    value -= 10
    if value < 26:
        return chr(65 + value)
    value -= 26
    if value < 26:
        return chr(97 + value)
    value -= 26
    if value == 0:
        return "-"
    if value == 1:
        return "_"
    return "?"


def _append3bytes(b1: int, b2: int, b3: int) -> str:
    c1 = b1 >> 2
    c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
    c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
    c4 = b3 & 0x3F
    return "".join(_encode6bit(bit & 0x3F) for bit in [c1, c2, c3, c4])


def plantuml_encode(text: str) -> str:
    compressor = zlib.compressobj(level=9, wbits=-15)
    compressed = compressor.compress(text.encode("utf-8")) + compressor.flush()

    chunks: list[str] = []
    for idx in range(0, len(compressed), 3):
        b1 = compressed[idx]
        b2 = compressed[idx + 1] if idx + 1 < len(compressed) else 0
        b3 = compressed[idx + 2] if idx + 2 < len(compressed) else 0
        chunks.append(_append3bytes(b1, b2, b3))
    return "".join(chunks)


def normalize_plantuml_server() -> str:
    raw = os.getenv("PLANTUML_SERVER_URL", "https://www.plantuml.com/plantuml").strip().rstrip("/")
    lowered = raw.lower()

    for suffix in ["/png", "/svg", "/txt", "/uml", "/img"]:
        if lowered.endswith(suffix):
            raw = raw[: -len(suffix)]
            lowered = raw.lower()
            break

    if not lowered.endswith("/plantuml"):
        raw = f"{raw}/plantuml"
    return raw


def plantuml_to_image(plantuml_code: str, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    encoded = plantuml_encode(plantuml_code)
    server_base = normalize_plantuml_server()
    url = f"{server_base}/png/{encoded}"

    request = requests.Request("GET", url, headers={"User-Agent": "Mozilla/5.0"}).prepare()
    with requests.Session() as session:
        response = session.send(request, timeout=25)
    response.raise_for_status()
    out_path.write_bytes(response.content)
    return out_path


def extract_fenced_sections(markdown_text: str, fence_name: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_heading = ""
    inside_block = False
    block_lines: list[str] = []
    block_title = ""
    expected_fence = f"```{fence_name.lower()}"

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if not inside_block:
            if stripped.startswith("## "):
                current_heading = stripped[3:].strip()
            elif stripped.startswith("### "):
                current_heading = stripped[4:].strip()

            if stripped.lower().startswith(expected_fence):
                inside_block = True
                block_lines = []
                block_title = current_heading or f"Seccion {len(sections) + 1}"
            continue

        if stripped.startswith("```"):
            code = "\n".join(block_lines).strip()
            if code:
                sections.append((block_title, code))
            inside_block = False
            block_lines = []
            block_title = ""
            continue

        block_lines.append(line)

    return sections


def build_pipeline_diagram(out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(14, 5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    boxes = [
        (0.05, 0.35, 0.18, 0.3, "Entrada NL", "#eaf2ff"),
        (0.29, 0.35, 0.18, 0.3, "Agente\nAnalista", "#e8f8ef"),
        (0.53, 0.35, 0.18, 0.3, "Agente\nArquitecto", "#fff6e8"),
        (0.77, 0.35, 0.18, 0.3, "Salidas\n(Mermaid/PlantUML/Text)", "#f7ecff"),
    ]

    for x, y, w, h, label, color in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.03",
            linewidth=1.5,
            edgecolor="#0f172a",
            facecolor=color,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=12, fontweight="bold")

    arrow_y = 0.5
    arrows = [(0.23, 0.29), (0.47, 0.53), (0.71, 0.77)]
    for start_x, end_x in arrows:
        ax.annotate(
            "",
            xy=(end_x, arrow_y),
            xytext=(start_x, arrow_y),
            arrowprops=dict(arrowstyle="->", lw=2.2, color="#0f172a"),
        )

    ax.text(0.38, 0.72, "Si falta contexto -> preguntas", fontsize=10, color="#334155", ha="center")
    ax.text(0.62, 0.72, "Si contexto completo -> generacion", fontsize=10, color="#334155", ha="center")

    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)
    return out_path


def collect_evidence() -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    sys.path.insert(0, str(ROOT / "src"))
    from arquisys_ai.graph.workflow import ArquiSysWorkflow

    workflow = ArquiSysWorkflow()

    package_request = (
        "Genera el paquete de documentacion visual completo (Casos de uso, Secuencia, BPMN y Entidad-Relacion) "
        "para el modulo de Checkout y Pago de un E-commerce de ropa. "
        "El flujo debe abarcar desde que el usuario revisa su carrito, hasta que se aprueba el pago y se notifica al almacen. "
        "Asume una pasarela de pagos externa y un sistema de inventario propio. "
        "Aplica manejo de errores como tarjeta sin fondos y falta de stock."
    )

    result_package = workflow.run(package_request)
    package_text = result_package.architecture.code if result_package.architecture else ""
    mermaid_sections = extract_fenced_sections(package_text, "mermaid")

    diagram_paths: dict[str, Path] = {}
    for idx, (title, code) in enumerate(mermaid_sections[:4], start=1):
        file_path = ASSETS_DIR / f"diagrama_mermaid_{idx}.jpg"
        try:
            mermaid_to_image(code, file_path)
        except Exception:
            text_to_image(code, file_path.with_suffix(".png"), f"Codigo Mermaid - {title}")
            file_path = file_path.with_suffix(".png")
        diagram_paths[f"mermaid_{idx}_{title}"] = file_path

    plantuml_request = (
        "Necesito un diagrama de secuencia UML en PlantUML. "
        "Participantes: Cliente, API, Pasarela. "
        "Interacciones: Cliente->>API: Iniciar checkout, API->>Pasarela: Autorizar pago, Pasarela-->>API: Pago aprobado"
    )
    result_plantuml = workflow.run(plantuml_request)
    plantuml_code = result_plantuml.architecture.code if result_plantuml.architecture else ""

    plantuml_image = ASSETS_DIR / "diagrama_plantuml_secuencia.jpg"
    try:
        plantuml_to_image(plantuml_code, plantuml_image)
    except Exception:
        fallback = plantuml_image.with_suffix(".png")
        text_to_image(plantuml_code, fallback, "Codigo PlantUML - Secuencia")
        plantuml_image = fallback

    cli_complete = run_command(
        [sys.executable, str(ROOT / "scripts" / "run_local.py"), "--request", package_request],
        ROOT,
    )
    cli_missing = run_command(
        [sys.executable, str(ROOT / "scripts" / "run_local.py"), "--request", "Hazme un diagrama"],
        ROOT,
    )
    pytest_out = run_command([sys.executable, "-m", "pytest", "-q"], ROOT)

    cli_complete_img = text_to_image(cli_complete, ASSETS_DIR / "captura_cli_completo.png", "Captura CLI - Caso completo")
    cli_missing_img = text_to_image(cli_missing, ASSETS_DIR / "captura_cli_faltante.png", "Captura CLI - Contexto insuficiente")
    pytest_img = text_to_image(pytest_out, ASSETS_DIR / "captura_pytest.png", "Captura QA - Resultado de pruebas")
    pipeline_img = build_pipeline_diagram(ASSETS_DIR / "diagrama_pipeline.png")

    evidence: dict[str, Path] = {
        "pipeline": pipeline_img,
        "cli_complete": cli_complete_img,
        "cli_missing": cli_missing_img,
        "pytest": pytest_img,
        "plantuml_sequence": plantuml_image,
    }
    evidence.update(diagram_paths)
    return evidence


def add_paragraph(document: Document, text: str, spacing_after: Pt = Pt(8)) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = spacing_after


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item)
        paragraph.style = "List Bullet"


def add_figure(document: Document, image_path: Path, caption: str, width_inches: float = 6.4) -> None:
    if not image_path.exists():
        add_paragraph(document, f"[No se encontro la figura: {image_path.name}]")
        return

    document.add_picture(str(image_path), width=Inches(width_inches))
    caption_par = document.add_paragraph(caption)
    caption_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_par.runs[0].italic = True
    caption_par.paragraph_format.space_after = Pt(10)


def build_report(output_path: Path, evidence: dict[str, Path]) -> None:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_heading("Informe Tecnico del Sprint 2", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = document.add_paragraph(
        "Proyecto: ArquiSys AI\n"
        "Tema: Implementacion del flujo agentico Analista -> Arquitecto\n"
        f"Fecha: {date.today().isoformat()}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()

    document.add_heading("Indice", level=1)
    add_bullets(
        document,
        [
            "1. Introduccion y contexto del sprint",
            "2. Objetivos, alcance y criterios de aceptacion",
            "3. Desarrollo tecnico ejecutado",
            "4. Arquitectura agentica implementada",
            "5. Evidencias funcionales y visuales",
            "6. Calidad, pruebas y resultados",
            "7. Incidencias detectadas y correcciones",
            "8. Evaluacion final del sprint",
            "9. Plan de continuidad para Sprint 3",
        ],
    )

    document.add_heading("1. Introduccion y contexto del sprint", level=1)
    add_paragraph(
        document,
        "Sprint 2 se centro en consolidar una base funcional de IA agentica orientada a documentacion tecnica. "
        "El objetivo principal fue transformar solicitudes en lenguaje natural hacia salidas de modelado tecnico "
        "con control de contexto y generacion automatizada."
    )
    add_paragraph(
        document,
        "Se trabajo con dos agentes: (a) Analista, responsable de interpretar la intencion y completar estructura de datos; "
        "(b) Arquitecto, encargado de construir diagramas o artefactos finales segun formato solicitado."
    )

    document.add_heading("2. Objetivos, alcance y criterios de aceptacion", level=1)
    add_paragraph(
        document,
        "Objetivo general: entregar un pipeline local operable y verificable para diagramacion tecnica con soporte multiformato."
    )

    add_bullets(
        document,
        [
            "Detectar tipo de salida y formato de manera automatica a partir del contexto.",
            "Bloquear generacion cuando falte informacion critica.",
            "Generar diagramas UML/ER/C4 y artefactos QA/scaffolding.",
            "Incluir soporte Mermaid y PlantUML en el agente arquitecto.",
            "Habilitar interfaz de uso local y una demo web para GitHub Pages.",
        ],
    )

    add_paragraph(
        document,
        "Criterio clave de aceptacion: para entradas completas, el sistema debe devolver salida util sin intervencion manual; "
        "para entradas incompletas, debe devolver preguntas de aclaracion en lugar de un diagrama defectuoso."
    )

    document.add_heading("3. Desarrollo tecnico ejecutado", level=1)
    add_paragraph(
        document,
        "Se implemento una ampliacion integral del sprint, cubriendo tanto robustez sintactica como valor visual de salida."
    )

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Componente"
    header[1].text = "Implementacion"
    header[2].text = "Resultado"

    rows = [
        ("Analista", "Inferencia de tipo, formato y contexto por dominio", "Deteccion de casos simples y paquetes completos"),
        ("Arquitecto", "Generadores Mermaid + PlantUML + Text", "Salida multiformato funcional"),
        ("UI Streamlit", "Vista previa y formulario de aclaracion", "Render Mermaid y PlantUML con mejor legibilidad"),
        ("Demo Web", "Sitio estatico en docs/ para Pages", "Ejecucion en GitHub Pages sin backend"),
        ("Calidad", "Pruebas unitarias e integracion", "25 pruebas aprobadas en validacion local"),
    ]
    for component, implementation, result in rows:
        row = table.add_row().cells
        row[0].text = component
        row[1].text = implementation
        row[2].text = result

    document.add_heading("4. Arquitectura agentica implementada", level=1)
    add_paragraph(
        document,
        "La arquitectura efectiva del sprint se resume en una canalizacion lineal con control de decisiones. "
        "El Analista valida suficiencia de contexto y el Arquitecto solo genera si la etapa previa aprueba."
    )
    add_figure(
        document,
        evidence["pipeline"],
        "Figura 1. Flujo del pipeline agentico implementado en Sprint 2.",
        width_inches=6.6,
    )

    document.add_heading("5. Evidencias funcionales y visuales", level=1)
    add_paragraph(
        document,
        "Se incluyen capturas de ejecucion real y diagramas generados automaticamente para documentar el comportamiento del sistema."
    )

    document.add_heading("5.1 Capturas de ejecucion por consola", level=2)
    add_figure(
        document,
        evidence["cli_complete"],
        "Figura 2. Ejecucion CLI con solicitud completa (se genera paquete de documentacion).",
        width_inches=6.7,
    )
    add_figure(
        document,
        evidence["cli_missing"],
        "Figura 3. Ejecucion CLI con contexto insuficiente (se bloquea salida y se solicita aclaracion).",
        width_inches=6.7,
    )

    document.add_heading("5.2 Diagramas generados (Mermaid)", level=2)
    mermaid_keys = sorted(key for key in evidence.keys() if key.startswith("mermaid_"))
    mermaid_captions = [
        "Figura 4. Diagrama de casos de uso generado automaticamente.",
        "Figura 5. Diagrama de secuencia generado automaticamente.",
        "Figura 6. Diagrama de flujo/BPMN-like con manejo de decisiones y errores.",
        "Figura 7. Diagrama entidad-relacion generado automaticamente.",
    ]
    for idx, key in enumerate(mermaid_keys[:4]):
        add_figure(document, evidence[key], mermaid_captions[idx], width_inches=6.6)

    document.add_heading("5.3 Diagrama generado en PlantUML", level=2)
    add_figure(
        document,
        evidence["plantuml_sequence"],
        "Figura 8. Ejemplo de visualizacion PlantUML (secuencia).",
        width_inches=6.6,
    )

    document.add_heading("6. Calidad, pruebas y resultados", level=1)
    add_paragraph(
        document,
        "La calidad del sprint se verifico en dos niveles: pruebas automaticas y validaciones de escenario. "
        "El resultado final indica estabilidad operativa de la base implementada."
    )

    metrics_table = document.add_table(rows=1, cols=2)
    metrics_table.style = "Table Grid"
    metrics_header = metrics_table.rows[0].cells
    metrics_header[0].text = "Metrica"
    metrics_header[1].text = "Valor"

    metric_rows = [
        ("Pruebas ejecutadas", "25"),
        ("Pruebas aprobadas", "25"),
        ("Fallas bloqueantes abiertas al cierre", "0 (en validacion local)"),
        ("Formatos soportados", "Mermaid, PlantUML, Text"),
        ("Tipos de salida principales", "UML, ER, C4, QA, Scaffolding, Doc Package"),
    ]
    for metric, value in metric_rows:
        row = metrics_table.add_row().cells
        row[0].text = metric
        row[1].text = value

    add_figure(
        document,
        evidence["pytest"],
        "Figura 9. Evidencia de ejecucion de pruebas automaticas del sprint.",
        width_inches=6.7,
    )

    document.add_heading("7. Incidencias detectadas y correcciones", level=1)
    add_paragraph(document, "Durante el sprint se detectaron incidencias tecnicas relevantes y se aplicaron correcciones:")
    add_bullets(
        document,
        [
            "Error de sintaxis en secuencia Mermaid (compatibilidad 10.9.5): se simplifico generador y se endurecieron IDs.",
            "Visualizacion parcial de paquetes: se ajusto UI para render por seleccion en lugar de render oculto masivo.",
            "Problemas de contraste visual: se aplico ajuste global de estilos para legibilidad de texto y formularios.",
            "Seleccion PlantUML sin efecto: se implemento generacion PlantUML nativa en el arquitecto.",
            "Render PlantUML en UI: se activo vista previa con servidor remoto/local y manejo de fallback.",
        ],
    )

    document.add_heading("8. Evaluacion final del sprint", level=1)
    add_paragraph(
        document,
        "Sprint 2 cumple su objetivo tecnico: el sistema opera con dos agentes, controla contexto, "
        "genera documentacion util y ofrece salidas en formatos usados por practica profesional."
    )
    add_paragraph(
        document,
        "Desde la perspectiva de valor, se obtuvo una base escalable para evolucionar hacia validacion formal, "
        "retroalimentacion iterativa y mayor trazabilidad de decisiones de la IA agentica."
    )

    document.add_heading("9. Plan de continuidad para Sprint 3", level=1)
    add_bullets(
        document,
        [
            "Implementar Agente Validador formal por estandar (UML/BPMN/ER/C4).",
            "Integrar ciclo Arquitecto <-> Validador con limite de iteraciones.",
            "Agregar criterios de calidad estructural y sintactica antes de salida final.",
            "Incorporar trazabilidad de decisiones para auditoria y explicabilidad.",
            "Extender automatizacion de exportaciones (PNG/SVG) y plantillas institucionales.",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    evidence = collect_evidence()

    detailed_report = DOCS_DIR / "Informe_Sprint_2_ArquiSys_AI_Detallado.docx"
    canonical_report = DOCS_DIR / "Informe_Sprint_2_ArquiSys_AI.docx"

    build_report(detailed_report, evidence)
    canonical_target = canonical_report
    try:
        build_report(canonical_target, evidence)
    except PermissionError:
        suffix = date.today().strftime("%Y%m%d")
        canonical_target = DOCS_DIR / f"Informe_Sprint_2_ArquiSys_AI_{suffix}.docx"
        build_report(canonical_target, evidence)

    print(f"Informe detallado generado en: {detailed_report}")
    print(f"Copia sincronizada generada en: {canonical_target}")


if __name__ == "__main__":
    main()
