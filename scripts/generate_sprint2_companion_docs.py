from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "report_assets"


def add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    paragraph.paragraph_format.space_after = Pt(8)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item)
        paragraph.style = "List Bullet"


def add_figure(document: Document, image_path: Path, caption: str, width_inches: float = 6.4) -> None:
    if not image_path.exists():
        add_paragraph(document, f"[Figura no encontrada: {image_path.name}]", italic=True)
        return

    document.add_picture(str(image_path), width=Inches(width_inches))
    caption_par = document.add_paragraph(caption)
    caption_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_par.runs[0].italic = True
    caption_par.paragraph_format.space_after = Pt(10)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value


def read_snippet(path: Path, start_line: int, end_line: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    start = max(1, start_line)
    end = min(len(lines), end_line)
    selected = lines[start - 1 : end]
    return "\n".join(selected)


def add_code_block(document: Document, title: str, code: str) -> None:
    document.add_heading(title, level=3)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(8)


def explain_code_line(line: str) -> str:
    stripped = line.strip()
    if not stripped:
        return "Linea en blanco para separar bloques logicos y mejorar legibilidad."
    if stripped.startswith("from ") or stripped.startswith("import "):
        return "Importa una dependencia necesaria para la ejecucion de este bloque."
    if stripped.startswith("def "):
        return "Declara una funcion; define la unidad de comportamiento reutilizable."
    if stripped.startswith("class "):
        return "Declara una clase que agrupa estado y metodos relacionados."
    if stripped.startswith("if "):
        return "Evalua una condicion para decidir el camino de ejecucion."
    if stripped.startswith("elif "):
        return "Evalua una condicion alternativa dentro de la misma decision."
    if stripped.startswith("else"):
        return "Define la ruta alternativa cuando condiciones previas no se cumplen."
    if stripped.startswith("for "):
        return "Itera sobre una coleccion para procesar elementos uno por uno."
    if stripped.startswith("while "):
        return "Itera mientras una condicion sea verdadera."
    if stripped.startswith("try"):
        return "Inicia bloque de control de errores para operaciones que pueden fallar."
    if stripped.startswith("except"):
        return "Captura y maneja una excepcion detectada en el bloque try."
    if stripped.startswith("with "):
        return "Abre un contexto controlado (recurso/archivo/sesion) con cierre automatico."
    if stripped.startswith("return ") or stripped == "return":
        return "Devuelve el resultado al nivel superior que llamo esta funcion."
    if stripped.startswith("raise "):
        return "Lanza una excepcion para reportar un error o estado invalido."
    if stripped.startswith("#"):
        return "Comentario de documentacion interna del codigo."
    if stripped in {"(", ")", "[", "]", "{", "}"}:
        return "Linea estructural que abre o cierra una expresion multilinea."
    if stripped.endswith(":"):
        return "Cabecera de bloque; la logica interna continua en lineas indentadas."
    if "=" in stripped and "==" not in stripped:
        return "Asigna o actualiza un valor que se utilizara en lineas posteriores."
    if stripped.startswith("@"):  # decorator
        return "Aplica un decorador para modificar comportamiento de funcion/clase."
    if stripped.startswith("}") or stripped.startswith("]") or stripped.startswith(")"):
        return "Cierra una estructura iniciada previamente."
    return "Instruccion operativa del bloque; aporta parte del flujo principal de ejecucion."


def add_fragment_with_line_comments(
    document: Document,
    title: str,
    file_path: Path,
    start_line: int,
    end_line: int,
    high_level_explanation: str,
) -> None:
    relative = file_path.relative_to(ROOT)
    lines = file_path.read_text(encoding="utf-8").splitlines()
    selected = lines[start_line - 1 : end_line]
    code = "\n".join(selected)

    add_code_block(document, title, code)
    add_paragraph(
        document,
        f"Archivo: {relative} | Lineas: {start_line}-{end_line}",
        bold=True,
    )
    add_paragraph(document, high_level_explanation)

    row_data: list[list[str]] = []
    for absolute_line, raw in enumerate(selected, start=start_line):
        code_cell = raw if raw.strip() else "(linea en blanco)"
        row_data.append([str(absolute_line), code_cell, explain_code_line(raw)])

    add_table(document, ["Linea", "Codigo", "Comentario linea por linea"], row_data)


def build_concepts_doc(output_path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_heading("Guia de conceptos y lectura del Informe Sprint 2", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = document.add_paragraph(
        "Proyecto: ArquiSys AI\n"
        f"Fecha: {date.today().isoformat()}\n"
        "Objetivo: explicar el informe principal a una audiencia no tecnica"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading("1. Como usar esta guia", level=1)
    add_paragraph(
        document,
        "Esta guia resume terminos tecnicos, clases de diagramas y significado de las figuras del informe de Sprint 2. "
        "Si una persona estudia este documento, puede explicar el informe principal sin necesidad de conocer programacion profunda."
    )

    document.add_heading("2. Palabras clave del proyecto", level=1)
    glossary_rows = [
        ["IA agentica", "Sistema que toma decisiones por etapas para cumplir una meta, no solo responde texto."],
        ["Agente Analista", "Componente que interpreta el requerimiento, extrae datos y detecta faltantes."],
        ["Agente Arquitecto", "Componente que genera la salida final (diagrama o artefacto textual)."],
        ["Contexto suficiente", "Informacion minima para construir un diagrama util sin inventar datos."],
        ["Pipeline", "Flujo secuencial Analista -> Arquitecto."],
        ["Mermaid", "Lenguaje textual para diagramas renderizados en web y markdown."],
        ["PlantUML", "Lenguaje textual alternativo para diagramas UML y otros modelos."],
        ["DOC_PACKAGE", "Salida que genera varios diagramas en una sola solicitud."],
        ["QA automatico", "Salida textual con escenarios base de prueba en formato tipo Gherkin."],
        ["Scaffolding", "Plantilla inicial de estructura de proyecto y endpoints base."],
        ["Fallback", "Ruta de respaldo cuando una visualizacion o formato falla."],
        ["Render", "Proceso de convertir codigo textual de diagrama en una imagen o grafico visual."],
    ]
    add_table(document, ["Termino", "Explicacion en lenguaje simple"], glossary_rows)

    document.add_heading("3. Tipos de diagramas que el sistema puede generar", level=1)
    diagram_rows = [
        [
            "Casos de uso (UML)",
            "Explicar funciones del sistema y quien interactua con ellas",
            "Analisis funcional inicial, presentaciones con negocio",
            "Actores, casos de uso",
        ],
        [
            "Secuencia (UML)",
            "Mostrar orden de mensajes entre componentes",
            "Diseño de APIs, validacion de flujos de negocio",
            "Participantes, interacciones",
        ],
        [
            "Flujo/BPMN-like",
            "Representar proceso paso a paso con decisiones",
            "Modelado de procesos operativos y manejo de errores",
            "Pasos de flujo",
        ],
        [
            "Entidad-Relacion (ER)",
            "Modelar estructura de datos y relaciones",
            "Diseño de base de datos",
            "Entidades, relaciones",
        ],
        [
            "C4 Context",
            "Mostrar sistema y actores externos",
            "Vista de alto nivel para arquitectura",
            "Personas y sistemas",
        ],
        [
            "C4 Container",
            "Mostrar bloques internos (web, api, db) y conexiones",
            "Diseño de componentes desplegables",
            "Sistemas, contenedores, relaciones",
        ],
    ]
    add_table(
        document,
        ["Tipo de salida", "Para que sirve", "Cuando usarlo", "Datos minimos"],
        diagram_rows,
    )

    document.add_heading("4. Explicacion de las figuras del informe principal", level=1)
    add_paragraph(
        document,
        "A continuacion se explica que muestra cada figura incluida en el informe detallado de Sprint 2."
    )

    add_figure(
        document,
        ASSETS_DIR / "diagrama_pipeline.png",
        "Figura A1. Pipeline agentico. Demuestra que el analista filtra el contexto antes de generar salida.",
        width_inches=6.6,
    )
    add_paragraph(
        document,
        "Lectura sugerida: primero entra la solicitud, luego el analista valida contexto, y solo despues el arquitecto genera diagramas."
    )

    add_figure(
        document,
        ASSETS_DIR / "captura_cli_completo.png",
        "Figura A2. Ejecucion con contexto completo. El sistema produce salida final.",
        width_inches=6.6,
    )
    add_paragraph(
        document,
        "Detalle de lectura de la Figura A2 (captura con mucho texto):",
        bold=True,
    )
    add_bullets(
        document,
        [
            "Bloque [ANALISIS]: es un JSON con la interpretacion del agente analista.",
            "Campo diagram_type: indica el tipo detectado (por ejemplo, doc_package o sequence).",
            "Campo diagram_format: define si la salida final sera Mermaid, PlantUML o Text.",
            "Campo missing_info: si esta vacio, el analista considera que el contexto es suficiente.",
            "Campo ready_for_architect: cuando vale true, habilita la etapa de generacion del Arquitecto.",
            "Bloque [SALIDA - ...]: contiene el codigo de los diagramas generados; si es paquete, incluye varias secciones.",
        ],
    )

    add_figure(
        document,
        ASSETS_DIR / "captura_cli_faltante.png",
        "Figura A3. Ejecucion con contexto incompleto. El sistema solicita aclaraciones.",
        width_inches=6.6,
    )
    add_paragraph(
        document,
        "Detalle de lectura de la Figura A3 (captura con mucho texto):",
        bold=True,
    )
    add_bullets(
        document,
        [
            "Campo missing_info: lista que dato falta (ejemplo: diagram_type, participants, entities).",
            "Campo clarifying_questions: preguntas concretas que el sistema devuelve para completar contexto.",
            "Mensaje [ESTADO]: confirma que no se genero diagrama por seguridad semantica.",
            "Interpretacion practica: la IA prioriza exactitud sobre generar una salida incompleta o engañosa.",
        ],
    )

    add_figure(
        document,
        ASSETS_DIR / "diagrama_mermaid_1.png",
        "Figura A4. Casos de uso para checkout/pago.",
        width_inches=6.4,
    )
    add_figure(
        document,
        ASSETS_DIR / "diagrama_mermaid_2.jpg",
        "Figura A5. Secuencia de mensajes en checkout/pago.",
        width_inches=6.4,
    )
    add_figure(
        document,
        ASSETS_DIR / "diagrama_mermaid_3.jpg",
        "Figura A6. Flujo de proceso con bifurcaciones de error.",
        width_inches=6.4,
    )
    add_figure(
        document,
        ASSETS_DIR / "diagrama_mermaid_4.jpg",
        "Figura A7. Modelo ER para datos de checkout/pago.",
        width_inches=6.4,
    )
    add_figure(
        document,
        ASSETS_DIR / "diagrama_plantuml_secuencia.jpg",
        "Figura A8. Ejemplo equivalente en PlantUML.",
        width_inches=6.4,
    )

    add_figure(
        document,
        ASSETS_DIR / "captura_pytest.png",
        "Figura A9. Captura de resultados de pruebas automaticas (texto tecnico).",
        width_inches=6.6,
    )
    add_paragraph(
        document,
        "Detalle de lectura de la Figura A9 (captura con mucho texto):",
        bold=True,
    )
    add_bullets(
        document,
        [
            "Linea collected: cantidad total de pruebas descubiertas por pytest.",
            "Marcador [100%]: confirma que se recorrio toda la suite planificada.",
            "Resumen passed: cantidad de pruebas aprobadas; evidencia estabilidad funcional del sprint.",
            "Uso en exposicion: demuestra control de calidad automatizado y no solo validacion manual.",
        ],
    )

    document.add_heading("5. Guion recomendado para exponer el informe", level=1)
    add_bullets(
        document,
        [
            "Paso 1: explicar el problema que resuelve ArquiSys AI (convertir texto a documentacion tecnica).",
            "Paso 2: presentar la arquitectura de dos agentes y la regla de contexto suficiente.",
            "Paso 3: mostrar figuras A4-A8 para demostrar tipos de salida y valor practico.",
            "Paso 4: destacar que se soporta Mermaid y PlantUML.",
            "Paso 5: usar la figura A9 para evidenciar calidad de pruebas (25/25).",
            "Paso 6: cerrar con pendientes de Sprint 3.",
        ],
    )

    document.add_heading("6. Preguntas frecuentes para audiencia no tecnica", level=1)
    faq_rows = [
        [
            "La IA inventa datos?",
            "No deberia. Si faltan datos, el analista devuelve preguntas en lugar de un diagrama final.",
        ],
        [
            "Se necesita entrenar un modelo desde cero?",
            "No para este sprint. Se usa orquestacion de agentes y reglas sobre componentes existentes.",
        ],
        [
            "Por que hay dos formatos (Mermaid y PlantUML)?",
            "Porque algunos equipos prefieren web/markdown (Mermaid) y otros UML clasico (PlantUML).",
        ],
        [
            "Que se valida en Sprint 3?",
            "Se incorporara un agente validador formal para revisar calidad sintactica y estructural.",
        ],
    ]
    add_table(document, ["Pregunta", "Respuesta recomendada"], faq_rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_technical_doc(output_path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_heading("Guia tecnica de librerias, codigo y bloques esenciales", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = document.add_paragraph(
        "Proyecto: ArquiSys AI\n"
        f"Fecha: {date.today().isoformat()}\n"
        "Objetivo: explicar como funciona internamente la solucion de Sprint 2"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading("1. Vista global de arquitectura", level=1)
    add_paragraph(
        document,
        "La aplicacion se organiza en capas: modelos de dominio, agentes, generadores por formato, "
        "workflow de orquestacion, scripts de UI y pruebas."
    )
    add_figure(
        document,
        ASSETS_DIR / "diagrama_pipeline.png",
        "Figura B1. Flujo base de ejecucion (Analista -> Arquitecto).",
        width_inches=6.5,
    )

    document.add_heading("2. Librerias usadas y rol dentro del sistema", level=1)
    libs_rows = [
        ["dataclasses / enum", "Modelado de resultados de analisis y arquitectura", "Estandar de Python"],
        ["streamlit", "Interfaz local interactiva", "Prototipo funcional rapido"],
        ["requests", "Consumo de servicios de render (Mermaid/PlantUML)", "Integracion HTTP simple"],
        ["pytest", "Validacion automatica de comportamiento", "Control de regresiones"],
        ["python-docx", "Generacion de informes en Word", "Entregables academicos/proyecto"],
        ["matplotlib + pillow", "Figuras e imagenes para informes", "Soporte visual de evidencia"],
        ["zlib/base64", "Codificacion PlantUML", "Construccion de URL de render"],
        ["mermaid.js + pako (docs web)", "Render de diagramas en GitHub Pages", "Demo sin backend"],
    ]
    add_table(document, ["Libreria", "Uso principal", "Motivo"], libs_rows)

    document.add_heading("3. Mapa de carpetas y responsabilidad", level=1)
    map_rows = [
        ["src/arquisys_ai/core/", "Enums y dataclasses del dominio (tipos, formatos, resultados)."],
        ["src/arquisys_ai/agents/", "Logica principal del Analista y Arquitecto."],
        ["src/arquisys_ai/generators/mermaid/", "Generadores Mermaid por tipo de diagrama."],
        ["src/arquisys_ai/generators/plantuml/", "Generadores PlantUML por tipo de diagrama."],
        ["src/arquisys_ai/generators/text/", "QA automatico y scaffolding textual."],
        ["src/arquisys_ai/graph/", "Workflow de ejecucion del pipeline."],
        ["scripts/", "Ejecucion local, UI y generadores de informes."],
        ["tests/", "Pruebas unitarias e integracion."],
        ["docs/", "Demo de GitHub Pages y entregables documentales."],
    ]
    add_table(document, ["Ruta", "Responsabilidad"], map_rows)

    document.add_heading("4. Bloques de codigo esenciales", level=1)
    add_paragraph(
        document,
        "En esta seccion se incluyen fragmentos reales y se explica por que cada bloque es critico para el comportamiento agentico."
    )

    workflow_path = ROOT / "src" / "arquisys_ai" / "graph" / "workflow.py"
    analyst_path = ROOT / "src" / "arquisys_ai" / "agents" / "analyst_agent.py"
    architect_path = ROOT / "src" / "arquisys_ai" / "agents" / "architect_agent.py"
    ui_path = ROOT / "scripts" / "ui_app.py"

    fragment_specs = [
        (
            "4.1 Orquestacion del pipeline (workflow.run)",
            workflow_path,
            13,
            20,
            "Bloque central que coordina el orden de ejecucion y evita generar salida cuando el analisis no esta listo.",
        ),
        (
            "4.2 Entrada del Analista (analyze)",
            analyst_path,
            23,
            34,
            "Bloque de entrada del analista: aplica heuristica, opcionalmente complementa con LLM y fija estado final del analisis.",
        ),
        (
            "4.3 Heuristica de extraccion de contexto",
            analyst_path,
            36,
            70,
            "Bloque donde se identifican etiquetas clave del lenguaje natural para convertir texto libre en estructura de datos.",
        ),
        (
            "4.4 Seleccion de generador por formato (ArchitectAgent.__init__)",
            architect_path,
            26,
            50,
            "Bloque de registro de generadores. Permite enrutar por tipo de diagrama y por formato de salida.",
        ),
        (
            "4.5 Build principal del Arquitecto",
            architect_path,
            51,
            107,
            "Bloque de decision principal del arquitecto: valida precondiciones, selecciona formato y construye la respuesta.",
        ),
        (
            "4.6 Generacion de paquete multidiagrama",
            architect_path,
            151,
            194,
            "Bloque encargado de ensamblar varios diagramas en una sola salida cuando se solicita doc_package.",
        ),
        (
            "4.7 Render Mermaid en UI",
            ui_path,
            188,
            251,
            "Bloque de interfaz que incrusta Mermaid.js y renderiza diagramas con estilos de lectura y manejo de errores.",
        ),
        (
            "4.8 Render PlantUML en UI",
            ui_path,
            309,
            327,
            "Bloque de interfaz para visualizacion PlantUML: codifica, consulta servidor y muestra imagen resultante.",
        ),
    ]

    for title_text, path, start, end, explanation in fragment_specs:
        add_fragment_with_line_comments(
            document=document,
            title=title_text,
            file_path=path,
            start_line=start,
            end_line=end,
            high_level_explanation=explanation,
        )

    document.add_heading("5. Bloques criticos por tipo de salida", level=1)
    critical_rows = [
        ["Mermaid Secuencia", "generators/mermaid/sequence_generator.py", "Compatibilidad con Mermaid 10.9.5 y parseo robusto"],
        ["Mermaid Casos de uso", "generators/mermaid/use_case_generator.py", "Incluye include/extend para mas valor semantico"],
        ["Mermaid Flujo", "generators/mermaid/flow_generator.py", "Agrega decisiones y ramas de error en checkout"],
        ["PlantUML Generators", "generators/plantuml/*.py", "Salida UML alternativa cuando el usuario pide PlantUML"],
        ["QA y Scaffolding", "generators/text/*.py", "Artefactos auxiliares para ciclo de desarrollo"],
    ]
    add_table(document, ["Bloque", "Ubicacion", "Por que es esencial"], critical_rows)

    document.add_heading("6. Estrategia de pruebas", level=1)
    add_paragraph(
        document,
        "Las pruebas estan divididas en unitarias (comportamiento local de agentes) e integracion (pipeline completo)."
    )
    add_figure(
        document,
        ASSETS_DIR / "captura_pytest.png",
        "Figura B2. Resultado de pruebas automaticas usadas para validar estabilidad.",
        width_inches=6.6,
    )
    add_bullets(
        document,
        [
            "tests/unit/test_analyst_agent.py: clasificacion, extraccion de contexto y preguntas.",
            "tests/unit/test_architect_agent.py: generacion por formato/tipo, incluyendo PlantUML.",
            "tests/integration/test_analyst_architect_flow.py: ejecucion extremo a extremo.",
        ],
    )

    document.add_heading("7. Si se quiere extender la IA agentica", level=1)
    add_bullets(
        document,
        [
            "Agregar Agente Validador que revise sintaxis y estructura antes de publicar salida.",
            "Agregar trazabilidad (log de decisiones) por cada etapa del pipeline.",
            "Agregar exportacion automatica a PNG/SVG y versionado de artefactos.",
            "Agregar adaptadores por dominio (salud, finanzas, retail) para mejorar precision inicial.",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    concepts_path = DOCS_DIR / "Guia_Explicativa_Sprint_2_Para_Exposicion.docx"
    technical_path = DOCS_DIR / "Guia_Tecnica_Librerias_y_Codigo_Sprint_2.docx"

    concepts_target = concepts_path
    technical_target = technical_path

    try:
        build_concepts_doc(concepts_target)
    except PermissionError:
        suffix = date.today().strftime("%Y%m%d")
        concepts_target = DOCS_DIR / f"Guia_Explicativa_Sprint_2_Para_Exposicion_{suffix}.docx"
        build_concepts_doc(concepts_target)

    try:
        build_technical_doc(technical_target)
    except PermissionError:
        suffix = date.today().strftime("%Y%m%d")
        technical_target = DOCS_DIR / f"Guia_Tecnica_Librerias_y_Codigo_Sprint_2_{suffix}.docx"
        build_technical_doc(technical_target)

    print(f"Guia de conceptos generada en: {concepts_target}")
    print(f"Guia tecnica generada en: {technical_target}")


if __name__ == "__main__":
    main()
